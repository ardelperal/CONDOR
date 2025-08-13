#!/usr/bin/env python3
"""
Script para convertir archivos DOCX a formato Markdown
Utiliza python-docx para extraer texto de archivos Word
"""

import os
import sys
import re
from pathlib import Path
from docx import Document

def clean_text(text):
    """Limpia y formatea el texto extraído del DOCX"""
    # Eliminar saltos de línea excesivos
    text = re.sub(r'\n\s*\n\s*\n', '\n\n', text)
    
    # Eliminar espacios al inicio y final de líneas
    lines = [line.strip() for line in text.split('\n')]
    
    # Reconstruir el texto
    cleaned_text = '\n'.join(lines)
    
    return cleaned_text

def extract_text_from_docx(docx_path):
    """Extrae texto de un archivo DOCX"""
    try:
        doc = Document(docx_path)
        text = ""
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # Detectar posibles títulos (texto en mayúsculas o con formato especial)
                para_text = paragraph.text.strip()
                
                # Si el párrafo está en mayúsculas y es corto, tratarlo como título
                if para_text.isupper() and len(para_text) < 100:
                    text += f"\n## {para_text}\n\n"
                else:
                    text += para_text + "\n\n"
        
        # Procesar tablas si las hay
        for table in doc.tables:
            text += "\n\n| "
            # Encabezados de tabla
            if table.rows:
                header_row = table.rows[0]
                headers = [cell.text.strip() for cell in header_row.cells]
                text += " | ".join(headers) + " |\n"
                text += "| " + " | ".join(["---" for _ in headers]) + " |\n"
                
                # Filas de datos
                for row in table.rows[1:]:
                    row_data = [cell.text.strip() for cell in row.cells]
                    text += "| " + " | ".join(row_data) + " |\n"
            text += "\n\n"
        
        return clean_text(text)
    
    except Exception as e:
        print(f"Error al procesar {docx_path}: {str(e)}")
        return None

def docx_to_markdown(docx_path, output_path=None):
    """Convierte un DOCX a Markdown"""
    if not os.path.exists(docx_path):
        print(f"Error: El archivo {docx_path} no existe")
        return False
    
    # Generar nombre de salida si no se proporciona
    if output_path is None:
        docx_name = Path(docx_path).stem
        output_path = f"{docx_name}.md"
    
    print(f"Convirtiendo {docx_path} a {output_path}...")
    
    # Extraer texto del DOCX
    text = extract_text_from_docx(docx_path)
    
    if text is None:
        return False
    
    # Crear encabezado Markdown
    docx_name = Path(docx_path).stem
    markdown_content = f"# {docx_name}\n\n"
    markdown_content += f"*Convertido automáticamente desde DOCX*\n\n"
    markdown_content += "---\n\n"
    markdown_content += text
    
    # Guardar archivo Markdown
    try:
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"✓ Conversión completada: {output_path}")
        return True
    
    except Exception as e:
        print(f"Error al guardar {output_path}: {str(e)}")
        return False

def main():
    """Función principal"""
    if len(sys.argv) < 2:
        print("Uso: python docx_to_markdown.py <archivo.docx> [salida.md]")
        sys.exit(1)
    
    docx_path = sys.argv[1]
    output_path = sys.argv[2] if len(sys.argv) > 2 else None
    
    success = docx_to_markdown(docx_path, output_path)
    
    if not success:
        sys.exit(1)

if __name__ == "__main__":
    main()